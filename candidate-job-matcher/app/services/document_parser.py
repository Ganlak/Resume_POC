"""
Document Parser Service
Extracts text from PDF, DOCX, and TXT files using LangChain
"""

import sys
from pathlib import Path

# Add project root to Python path
project_root = Path(__file__).parent.parent.parent
sys.path.insert(0, str(project_root))

import logging
from typing import Tuple, Optional, Dict

# LangChain Document Loaders
from langchain_community.document_loaders import (
    PyPDFLoader,
    Docx2txtLoader,
    TextLoader
)

# Validation
from app.utils.validators import (
    validate_uploaded_file,
    validate_text_content,
    extract_email,
    extract_phone,
    extract_name_from_filename,
    extract_name_from_text
)
from app.utils.helpers import clean_text

logger = logging.getLogger(__name__)


# ==========================================
# LangChain-based Document Parsers
# ==========================================

def parse_pdf_langchain(file_path: str) -> Tuple[bool, str, str]:
    """
    Parse PDF using LangChain PyPDFLoader
    
    Args:
        file_path: Path to PDF file
        
    Returns:
        Tuple of (success, text_content, error_message)
    """
    try:
        loader = PyPDFLoader(file_path)
        documents = loader.load()
        
        # Combine all pages
        text_content = "\n\n".join([doc.page_content for doc in documents])
        
        if not text_content.strip():
            return False, "", "No text content found in PDF"
        
        logger.info(f"Successfully parsed PDF with LangChain: {file_path}")
        return True, text_content, ""
    
    except Exception as e:
        logger.error(f"Error parsing PDF with LangChain: {e}")
        # Try fallback with PyPDF2 directly
        return parse_pdf_fallback(file_path)


def parse_pdf_fallback(file_path: str) -> Tuple[bool, str, str]:
    """
    Fallback PDF parser using PyPDF2 directly
    
    Args:
        file_path: Path to PDF file
        
    Returns:
        Tuple of (success, text_content, error_message)
    """
    try:
        from PyPDF2 import PdfReader
        
        logger.info(f"Using PyPDF2 fallback for: {file_path}")
        reader = PdfReader(file_path)
        text_content = []
        
        for page_num, page in enumerate(reader.pages, 1):
            try:
                text = page.extract_text()
                if text.strip():
                    text_content.append(text)
                    logger.debug(f"Extracted text from page {page_num}")
            except Exception as e:
                logger.warning(f"Could not extract text from page {page_num}: {e}")
                continue
        
        full_text = "\n\n".join(text_content)
        
        if not full_text.strip():
            return False, "", "No text content could be extracted from PDF"
        
        logger.info(f"Successfully parsed PDF with PyPDF2 fallback: {file_path}")
        return True, full_text, ""
        
    except Exception as e:
        logger.error(f"PyPDF2 fallback also failed: {e}")
        return False, "", f"Failed to parse PDF: {str(e)}"


def parse_docx_langchain(file_path: str) -> Tuple[bool, str, str]:
    """
    Parse DOCX using LangChain Docx2txtLoader
    
    Args:
        file_path: Path to DOCX file
        
    Returns:
        Tuple of (success, text_content, error_message)
    """
    try:
        loader = Docx2txtLoader(file_path)
        documents = loader.load()
        
        # Combine all document parts
        text_content = "\n\n".join([doc.page_content for doc in documents])
        
        if not text_content.strip():
            return False, "", "No text content found in DOCX"
        
        logger.info(f"Successfully parsed DOCX with LangChain: {file_path}")
        return True, text_content, ""
    
    except Exception as e:
        logger.error(f"Error parsing DOCX with LangChain: {e}")
        # Try fallback with python-docx directly
        return parse_docx_fallback(file_path)


def parse_docx_fallback(file_path: str) -> Tuple[bool, str, str]:
    """
    Fallback DOCX parser using python-docx directly
    
    Args:
        file_path: Path to DOCX file
        
    Returns:
        Tuple of (success, text_content, error_message)
    """
    try:
        from docx import Document
        
        logger.info(f"Using python-docx fallback for: {file_path}")
        doc = Document(file_path)
        
        # Extract all paragraphs
        paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]
        
        # Extract text from tables
        table_text = []
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join([cell.text.strip() for cell in row.cells if cell.text.strip()])
                if row_text:
                    table_text.append(row_text)
        
        # Combine all text
        all_text = paragraphs + table_text
        full_text = "\n".join(all_text)
        
        if not full_text.strip():
            return False, "", "No text content could be extracted from DOCX"
        
        logger.info(f"Successfully parsed DOCX with python-docx fallback: {file_path}")
        return True, full_text, ""
        
    except Exception as e:
        logger.error(f"python-docx fallback also failed: {e}")
        return False, "", f"Failed to parse DOCX: {str(e)}"


def parse_txt_langchain(file_path: str) -> Tuple[bool, str, str]:
    """
    Parse TXT using LangChain TextLoader
    
    Args:
        file_path: Path to TXT file
        
    Returns:
        Tuple of (success, text_content, error_message)
    """
    try:
        # Try different encodings
        encodings = ['utf-8', 'latin-1', 'cp1252', 'iso-8859-1']
        
        for encoding in encodings:
            try:
                loader = TextLoader(file_path, encoding=encoding)
                documents = loader.load()
                
                text_content = "\n\n".join([doc.page_content for doc in documents])
                
                if text_content.strip():
                    logger.info(f"Successfully parsed TXT with LangChain ({encoding}): {file_path}")
                    return True, text_content, ""
            
            except Exception:
                continue
        
        return False, "", "Failed to parse TXT with supported encodings"
    
    except Exception as e:
        logger.error(f"Error parsing TXT with LangChain: {e}")
        return False, "", str(e)


# ==========================================
# Main Document Parser
# ==========================================

def parse_document(file_path: str, file_type: str) -> Tuple[bool, str, str]:
    """
    Parse document based on file type using LangChain
    
    Args:
        file_path: Path to file
        file_type: File extension (pdf, docx, txt)
        
    Returns:
        Tuple of (success, text_content, error_message)
    """
    file_type = file_type.lower().strip('.')
    
    # Parse based on file type (with built-in fallbacks)
    if file_type == 'pdf':
        return parse_pdf_langchain(file_path)
    elif file_type == 'docx':
        return parse_docx_langchain(file_path)
    elif file_type == 'txt':
        return parse_txt_langchain(file_path)
    else:
        return False, "", f"Unsupported file type: {file_type}"


# ==========================================
# Resume Information Extractor
# ==========================================

def extract_candidate_info(text: str, filename: str) -> Dict[str, Optional[str]]:
    """
    Extract candidate information from resume text
    
    Args:
        text: Resume text content
        filename: Original filename
        
    Returns:
        Dictionary with extracted information
    """
    info = {
        "name": None,
        "email": None,
        "phone": None
    }
    
    # Extract email
    info["email"] = extract_email(text)
    
    # Extract phone
    info["phone"] = extract_phone(text)
    
    # Extract name (try from text first, then filename)
    info["name"] = extract_name_from_text(text)
    
    if not info["name"]:
        info["name"] = extract_name_from_filename(filename)
    
    # If still no name, use filename without extension
    if not info["name"]:
        info["name"] = Path(filename).stem.replace('_', ' ').replace('-', ' ').title()
    
    return info


# ==========================================
# Complete Resume Parser
# ==========================================

def parse_resume(
    file_path: str,
    filename: str,
    file_size: int
) -> Dict[str, any]:
    """
    Complete resume parsing with validation and extraction using LangChain
    
    Args:
        file_path: Path to resume file
        filename: Original filename
        file_size: File size in bytes
        
    Returns:
        Dictionary with parsing results
    """
    result = {
        "success": False,
        "text": "",
        "cleaned_text": "",
        "candidate_info": {},
        "error": None,
        "file_info": {
            "filename": filename,
            "file_size": file_size,
            "file_type": Path(filename).suffix.lower().strip('.')
        }
    }
    
    # Validate file
    is_valid, message = validate_uploaded_file(filename, file_size)
    if not is_valid:
        result["error"] = message
        return result
    
    # Parse document using LangChain
    file_type = result["file_info"]["file_type"]
    success, text, error = parse_document(file_path, file_type)
    
    if not success:
        result["error"] = error
        return result
    
    result["text"] = text
    
    # Clean text
    cleaned = clean_text(text)
    result["cleaned_text"] = cleaned
    
    # Validate text content
    is_valid, message = validate_text_content(cleaned, min_length=50)
    if not is_valid:
        result["error"] = message
        return result
    
    # Extract candidate information
    candidate_info = extract_candidate_info(cleaned, filename)
    result["candidate_info"] = candidate_info
    
    result["success"] = True
    
    logger.info(f"Successfully parsed resume: {filename}")
    logger.info(f"  - Name: {candidate_info.get('name')}")
    logger.info(f"  - Email: {candidate_info.get('email')}")
    logger.info(f"  - Phone: {candidate_info.get('phone')}")
    logger.info(f"  - Text length: {len(cleaned)} characters")
    
    return result


# ==========================================
# Batch Processing
# ==========================================

def parse_multiple_resumes(file_paths: list) -> list:
    """
    Parse multiple resumes in batch
    
    Args:
        file_paths: List of tuples (file_path, filename, file_size)
        
    Returns:
        List of parsing results
    """
    results = []
    
    for file_path, filename, file_size in file_paths:
        result = parse_resume(file_path, filename, file_size)
        results.append(result)
    
    success_count = sum(1 for r in results if r["success"])
    logger.info(f"Batch parsing complete: {success_count}/{len(results)} successful")
    
    return results


# ==========================================
# Test Functions
# ==========================================

if __name__ == "__main__":
    """Test document parser with LangChain"""
    
    print("=" * 60)
    print("LangChain Document Parser Test")
    print("=" * 60)
    
    # Create test text file
    test_file_path = Path("test_resume.txt")
    test_content = """
    JOHN DOE
    Software Engineer
    
    Email: john.doe@example.com
    Phone: +1-123-456-7890
    
    EXPERIENCE
    Senior Software Engineer at Tech Corp
    - Developed web applications using Python and React
    - Led team of 5 developers
    - Implemented CI/CD pipelines
    
    SKILLS
    Python, JavaScript, React, SQL, Docker, AWS
    
    EDUCATION
    BS Computer Science, MIT, 2015
    """
    
    # Write test file
    with open(test_file_path, 'w') as f:
        f.write(test_content)
    
    print("\nTest File Created: test_resume.txt")
    print("\nParsing Resume with LangChain...")
    
    # Parse resume
    result = parse_resume(
        file_path=str(test_file_path),
        filename="test_resume.txt",
        file_size=test_file_path.stat().st_size
    )
    
    # Display results
    print("\n" + "=" * 60)
    print("PARSING RESULTS")
    print("=" * 60)
    print(f"Success: {result['success']}")
    print(f"Error: {result['error']}")
    print(f"\nCandidate Information:")
    print(f"  Name: {result['candidate_info'].get('name')}")
    print(f"  Email: {result['candidate_info'].get('email')}")
    print(f"  Phone: {result['candidate_info'].get('phone')}")
    print(f"\nText Length: {len(result['cleaned_text'])} characters")
    print(f"\nFirst 200 characters:")
    print(result['cleaned_text'][:200])
    
    # Cleanup
    test_file_path.unlink()
    print("\n" + "=" * 60)
    print("Test file cleaned up")
    print("LangChain parsing successful!")